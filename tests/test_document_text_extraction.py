# Tests app/ingestion/document_text_extraction.py: the PDF/DOCX parsing
# shared by the Google Drive and SharePoint connectors. No network.
from io import BytesIO

from app.ingestion.document_text_extraction import extract_docx_text, extract_pdf_text


def test_extract_pdf_text_returns_empty_for_a_blank_page():
    # A page with no text layer (e.g. a scanned/image-only PDF) should
    # produce "" rather than raise: the caller treats that as "skip this
    # file", not an error.
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)

    assert extract_pdf_text(buf.getvalue()) == ""


def test_extract_docx_text_extracts_real_paragraphs():
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = BytesIO()
    doc.save(buf)

    text = extract_docx_text(buf.getvalue())
    assert "First paragraph." in text
    assert "Second paragraph." in text
