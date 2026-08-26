# GoogleDriveConnector tests -- no real network, no real Google credentials.
# Auth (service_account.Credentials.from_service_account_info/.refresh) and
# the Drive REST calls (httpx.AsyncClient) are both monkeypatched, so these
# run free and fast, same spirit as test_connectors.py's web connector tests.
import asyncio
import json
from io import BytesIO

import httpx
import pytest

from app.ingestion.connector_base import ConnectorFetchError
from app.ingestion.google_drive_source import (
    GoogleDriveConnector,
    _extract_docx_text,
    _extract_folder_id,
    _extract_pdf_text,
)

_FAKE_SERVICE_ACCOUNT_JSON = json.dumps(
    {
        "type": "service_account",
        "project_id": "test-project",
        "private_key_id": "abc123",
        "private_key": "-----BEGIN PRIVATE KEY-----\nfake\n-----END PRIVATE KEY-----\n",
        "client_email": "demo@test-project.iam.gserviceaccount.com",
        "token_uri": "https://oauth2.googleapis.com/token",
    }
)


def test_extract_folder_id_from_bare_id():
    assert _extract_folder_id("1AbC-defGHI_23") == "1AbC-defGHI_23"


def test_extract_folder_id_from_folder_url():
    url = "https://drive.google.com/drive/folders/1AbC-defGHI_23?usp=sharing"
    assert _extract_folder_id(url) == "1AbC-defGHI_23"


def test_extract_folder_id_from_open_id_url():
    url = "https://drive.google.com/open?id=1AbC-defGHI_23"
    assert _extract_folder_id(url) == "1AbC-defGHI_23"


def test_extract_folder_id_rejects_garbage():
    with pytest.raises(ConnectorFetchError):
        _extract_folder_id("not a folder id or link!! ")


def test_source_description_includes_folder_id():
    connector = GoogleDriveConnector("1AbC-defGHI_23")
    assert "1AbC-defGHI_23" in connector.source_description()


def test_content_hash_matches_shared_hash_records_helper():
    from app.ingestion.connector_base import hash_records
    from app.ingestion.file_source import SourceRecord

    records = [SourceRecord(name="a", body="x", source_description="d")]
    connector = GoogleDriveConnector("1AbC-defGHI_23")
    assert connector.content_hash(records) == hash_records(records)


def test_fetch_fails_clearly_when_not_configured(monkeypatch):
    monkeypatch.setattr("app.ingestion.google_drive_source.settings.google_drive_service_account_json", "")
    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="isn't configured"):
        asyncio.run(connector.fetch())


def test_fetch_fails_clearly_on_malformed_credentials(monkeypatch):
    monkeypatch.setattr(
        "app.ingestion.google_drive_source.settings.google_drive_service_account_json", "{not valid json"
    )
    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="misconfigured"):
        asyncio.run(connector.fetch())


class _FakeCredentials:
    def __init__(self, *a, **kw):
        self.token = None

    def refresh(self, request):
        self.token = "fake-access-token"


def _patch_auth(monkeypatch):
    monkeypatch.setattr(
        "app.ingestion.google_drive_source.settings.google_drive_service_account_json",
        _FAKE_SERVICE_ACCOUNT_JSON,
    )
    monkeypatch.setattr(
        "app.ingestion.google_drive_source.service_account.Credentials.from_service_account_info",
        lambda info, scopes: _FakeCredentials(),
    )


class _FakeResponse:
    def __init__(self, status_code=200, json_data=None, text=""):
        self.status_code = status_code
        self._json_data = json_data
        self.text = text

    def json(self):
        return self._json_data

    def raise_for_status(self):
        if self.status_code >= 400:
            raise httpx.HTTPStatusError("error", request=None, response=self)


def test_fetch_reads_supported_files_and_skips_the_rest(monkeypatch):
    _patch_auth(monkeypatch)

    files = [
        {"id": "doc1", "name": "Notes", "mimeType": "application/vnd.google-apps.document"},
        {"id": "sheet1", "name": "Budget", "mimeType": "application/vnd.google-apps.spreadsheet"},
        {"id": "slides1", "name": "Pitch", "mimeType": "application/vnd.google-apps.presentation"},
        {"id": "txt1", "name": "readme.txt", "mimeType": "text/plain"},
        {"id": "img1", "name": "photo.png", "mimeType": "image/png"},
        {"id": "sub1", "name": "Subfolder", "mimeType": "application/vnd.google-apps.folder"},
    ]

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            if url.endswith("/files"):
                return _FakeResponse(200, json_data={"files": files})
            if url.endswith("/doc1/export"):
                assert params["mimeType"] == "text/plain"
                return _FakeResponse(200, text="Doc content here.")
            if url.endswith("/sheet1/export"):
                assert params["mimeType"] == "text/csv"
                return _FakeResponse(200, text="col1,col2\n1,2")
            if url.endswith("/slides1/export"):
                assert params["mimeType"] == "text/plain"
                return _FakeResponse(200, text="Slide content here.")
            if url.endswith("/txt1"):
                assert params["alt"] == "media"
                return _FakeResponse(200, text="Plain text content.")
            raise AssertionError(f"unexpected URL {url}")

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    records = asyncio.run(connector.fetch())

    assert len(records) == 4
    bodies = {r.body for r in records}
    assert "Doc content here." in bodies
    assert "col1,col2\n1,2" in bodies
    assert "Slide content here." in bodies
    assert "Plain text content." in bodies
    names = {r.source_description for r in records}
    assert "Google Drive (Notes)" in names
    assert "Google Drive (Budget)" in names
    assert "Google Drive (Pitch)" in names
    assert "Google Drive (readme.txt)" in names


def test_fetch_raises_clear_error_on_403(monkeypatch):
    _patch_auth(monkeypatch)

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            return _FakeResponse(403)

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="Access denied"):
        asyncio.run(connector.fetch())


def test_fetch_raises_when_no_supported_files(monkeypatch):
    _patch_auth(monkeypatch)

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            return _FakeResponse(200, json_data={"files": [{"id": "x", "name": "a.png", "mimeType": "image/png"}]})

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="No supported files"):
        asyncio.run(connector.fetch())


def test_extract_pdf_text_returns_empty_for_a_blank_page():
    # A page with no text layer (e.g. a scanned/image-only PDF) should
    # produce "" rather than raise -- the caller treats that as "skip this
    # file", not an error.
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)

    assert _extract_pdf_text(buf.getvalue()) == ""


def test_extract_docx_text_extracts_real_paragraphs():
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = BytesIO()
    doc.save(buf)

    text = _extract_docx_text(buf.getvalue())
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_fetch_reads_pdf_and_docx_via_the_binary_parsers(monkeypatch):
    _patch_auth(monkeypatch)

    files = [
        {"id": "pdf1", "name": "report.pdf", "mimeType": "application/pdf"},
        {
            "id": "docx1",
            "name": "memo.docx",
            "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    ]

    class _FakeBinaryResponse:
        def __init__(self, content: bytes):
            self.status_code = 200
            self.content = content

        def raise_for_status(self):
            pass

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            if url.endswith("/files"):
                return _FakeResponse(200, json_data={"files": files})
            if url.endswith("/pdf1"):
                return _FakeBinaryResponse(b"fake-pdf-bytes")
            if url.endswith("/docx1"):
                return _FakeBinaryResponse(b"fake-docx-bytes")
            raise AssertionError(f"unexpected URL {url}")

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())
    monkeypatch.setattr(
        "app.ingestion.google_drive_source._BINARY_PARSERS",
        {
            "application/pdf": lambda data: f"parsed:{data.decode()}",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": lambda data: f"parsed:{data.decode()}",
        },
    )

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    records = asyncio.run(connector.fetch())

    assert len(records) == 2
    bodies = {r.body for r in records}
    assert "parsed:fake-pdf-bytes" in bodies
    assert "parsed:fake-docx-bytes" in bodies


def test_fetch_skips_a_file_whose_parser_raises(monkeypatch):
    _patch_auth(monkeypatch)

    files = [{"id": "bad1", "name": "corrupt.pdf", "mimeType": "application/pdf"}]

    class _FakeBinaryResponse:
        status_code = 200
        content = b"not a real pdf"

        def raise_for_status(self):
            pass

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            if url.endswith("/files"):
                return _FakeResponse(200, json_data={"files": files})
            return _FakeBinaryResponse()

    def _raising_parser(data: bytes) -> str:
        raise ValueError("corrupt PDF")

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())
    monkeypatch.setattr(
        "app.ingestion.google_drive_source._BINARY_PARSERS", {"application/pdf": _raising_parser}
    )

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="No supported files"):
        asyncio.run(connector.fetch())


def test_fetch_skips_a_binary_file_over_the_size_cap(monkeypatch):
    _patch_auth(monkeypatch)

    files = [{"id": "big1", "name": "huge.pdf", "mimeType": "application/pdf"}]

    class _FakeBinaryResponse:
        status_code = 200
        content = b"x" * (16 * 1024 * 1024)  # over _MAX_BINARY_BYTES (15MB)

        def raise_for_status(self):
            pass

    class _FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def get(self, url, headers=None, params=None):
            if url.endswith("/files"):
                return _FakeResponse(200, json_data={"files": files})
            return _FakeBinaryResponse()

    called = False

    def _tracking_parser(data: bytes) -> str:
        nonlocal called
        called = True
        return "should not run"

    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: _FakeClient())
    monkeypatch.setattr(
        "app.ingestion.google_drive_source._BINARY_PARSERS", {"application/pdf": _tracking_parser}
    )

    connector = GoogleDriveConnector("1AbC-defGHI_23")
    with pytest.raises(ConnectorFetchError, match="No supported files"):
        asyncio.run(connector.fetch())
    assert called is False
